"""
Enhanced File parsing utilities for resume processing
Uses PyMuPDF (fitz) for better PDF extraction and includes formatting checks
"""

import fitz  # PyMuPDF
from docx import Document
from typing import Dict, Any, List
import io
import re
from PIL import Image

class FileParser:
    """
    Enhanced file parser with better PDF extraction and formatting analysis
    """
    
    def __init__(self):
        """Initialize parser with supported formats"""
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
    
    def parse_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Main method to parse any supported file type with enhanced extraction
        
        Args:
            file_content: The file data as bytes
            filename: The name of the file
            
        Returns:
            Dictionary with parsed content, metadata, and formatting analysis
        """
        file_extension = filename.lower().split('.')[-1]
        
        if file_extension == 'pdf':
            return self._parse_pdf_enhanced(file_content)
        elif file_extension in ['docx', 'doc']:
            return self._parse_docx_enhanced(file_content)
        elif file_extension == 'txt':
            return self._parse_txt(file_content)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _parse_pdf_enhanced(self, file_content: bytes) -> Dict[str, Any]:
        """
        Enhanced PDF parsing using PyMuPDF (fitz) for better text extraction
        Also includes formatting and image detection
        """
        try:
            pdf_file = io.BytesIO(file_content)
            doc = fitz.open(stream=pdf_file, filetype="pdf")
            
            text = ""
            images_count = 0
            tables_detected = False
            fonts_used = set()
            formatting_issues = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Method 1: Standard text extraction
                page_text = page.get_text()
                
                # Method 2: Extract text blocks with coordinates for better structure
                blocks = page.get_text("blocks")
                
                # Method 3: Extract text with layout preservation
                text_dict = page.get_text("dict")
                
                # Combine extraction methods for better content
                combined_text = ""
                
                # Process blocks to maintain structure and get more content
                for block in blocks:
                    if len(block) >= 5:  # Valid text block
                        block_text = block[4].strip()
                        if block_text:
                            combined_text += block_text + "\n"
                
                # Use the better extraction method
                if combined_text.strip():
                    text += combined_text + "\n"
                else:
                    text += page_text + "\n"
                
                # Detect images (ATS red flag)
                images = page.get_images()
                images_count += len(images)
                
                # Detect fonts (unusual fonts can cause ATS issues)
                for block in text_dict.get("blocks", []):
                    if "lines" in block:
                        for line in block["lines"]:
                            for span in line["spans"]:
                                fonts_used.add(span.get("font", "Unknown"))
                
                # Detect tables (can cause parsing issues)
                if self._detect_tables_in_text(page_text):
                    tables_detected = True
                
                # Detect text formatting issues
                if self._detect_formatting_issues(page_text):
                    formatting_issues.append(f"Page {page_num + 1}: Potential formatting issues detected")
            
            # Analyze formatting issues
            if images_count > 0:
                formatting_issues.append(f"Contains {images_count} image(s) - may cause ATS parsing issues")
            
            if tables_detected:
                formatting_issues.append("Tables detected - ensure content is also in plain text")
            
            if len(fonts_used) > 3:
                formatting_issues.append(f"Multiple fonts detected ({len(fonts_used)}) - stick to 1-2 standard fonts")
            
            # Check for unusual fonts
            common_fonts = ['Times', 'Arial', 'Helvetica', 'Calibri', 'Georgia']
            unusual_fonts = [f for f in fonts_used if not any(cf in f for cf in common_fonts)]
            if unusual_fonts:
                formatting_issues.append(f"Unusual fonts detected: {', '.join(list(unusual_fonts)[:3])}")
            
            # Store page count before closing
            page_count = len(doc)
            doc.close()
            
            # Calculate statistics
            word_count = len(text.split())
            character_count = len(text)
            
            return {
                "text": text.strip(),
                "word_count": word_count,
                "character_count": character_count,
                "page_count": page_count,
                "file_type": "pdf",
                "formatting_analysis": {
                    "images_count": images_count,
                    "tables_detected": tables_detected,
                    "fonts_count": len(fonts_used),
                    "fonts_used": list(fonts_used),
                    "formatting_issues": formatting_issues,
                    "ats_friendly": len(formatting_issues) == 0
                }
            }
            
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")
    
    def _parse_docx_enhanced(self, file_content: bytes) -> Dict[str, Any]:
        """
        Enhanced DOCX parsing with formatting analysis
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text = ""
            tables_count = len(doc.tables)
            images_count = 0
            formatting_issues = []
            
            # Extract text from paragraphs with better structure preservation
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    text += para_text + "\n"
            
            # Extract text from tables with better formatting
            for table in doc.tables:
                table_text = ""
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text += " | ".join(row_text) + "\n"
                if table_text:
                    text += table_text + "\n"
            
            # Extract text from headers and footers
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        header_text = paragraph.text.strip()
                        if header_text:
                            text += header_text + "\n"
                
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        footer_text = paragraph.text.strip()
                        if footer_text:
                            text += footer_text + "\n"
            
            # Count images/shapes
            try:
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        images_count += 1
            except:
                pass
            
            # Analyze formatting issues
            if images_count > 0:
                formatting_issues.append(f"Contains {images_count} image(s) - may cause ATS issues")
            
            if tables_count > 2:
                formatting_issues.append(f"Contains {tables_count} tables - consider simplifying layout")
            
            # Calculate statistics
            word_count = len(text.split())
            character_count = len(text)
            
            return {
                "text": text.strip(),
                "word_count": word_count,
                "character_count": character_count,
                "paragraph_count": len(doc.paragraphs),
                "file_type": "docx",
                "formatting_analysis": {
                    "images_count": images_count,
                    "tables_count": tables_count,
                    "tables_detected": tables_count > 0,
                    "formatting_issues": formatting_issues,
                    "ats_friendly": len(formatting_issues) == 0
                }
            }
            
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")
    
    def _parse_txt(self, file_content: bytes) -> Dict[str, Any]:
        """
        Parse TXT files (always ATS-friendly)
        """
        try:
            text = file_content.decode('utf-8')
            word_count = len(text.split())
            character_count = len(text)
            line_count = len(text.split('\n'))
            
            return {
                "text": text.strip(),
                "word_count": word_count,
                "character_count": character_count,
                "line_count": line_count,
                "file_type": "txt",
                "formatting_analysis": {
                    "images_count": 0,
                    "tables_detected": False,
                    "formatting_issues": [],
                    "ats_friendly": True
                }
            }
            
        except Exception as e:
            raise Exception(f"Error parsing TXT: {str(e)}")
    
    def _detect_tables_in_text(self, text: str) -> bool:
        """
        Simple table detection based on text patterns
        """
        lines = text.split('\n')
        
        # Look for patterns like multiple tabs or spaces indicating columns
        tabular_lines = 0
        for line in lines:
            if '\t' in line or re.search(r'\s{3,}', line):
                tabular_lines += 1
        
        # If more than 3 lines look like tables, flag it
        return tabular_lines > 3
    
    def _detect_formatting_issues(self, text: str) -> bool:
        """Detect potential formatting issues in text"""
        # Check for excessive special characters
        special_char_count = len(re.findall(r'[^\w\s\.\,\;\:\!\?\-\(\)]', text))
        if special_char_count > len(text) * 0.1:  # More than 10% special characters
            return True
        
        # Check for excessive whitespace
        if re.search(r'\s{5,}', text):  # 5 or more consecutive spaces
            return True
        
        # Check for mixed case issues
        if re.search(r'[A-Z]{5,}', text):  # 5 or more consecutive uppercase letters
            return True
        
        return False

# Create global instance
file_parser = FileParser()
